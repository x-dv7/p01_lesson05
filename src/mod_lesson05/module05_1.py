"""Основной модуль приложения lesson05."""

from pathlib import Path
from datetime import datetime
import tempfile
import zipfile
# import chardet
import json
import csv
from openpyxl import Workbook
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


def  scan_directory(path, base_path, archive_prefix=""):
    """
    Рекурсивно сканирует директорию и возвращает список словарей с информацией о файлах и папках.
    """
    entries = []
    path = Path(path)

    # Сортируем по относительному пути, чтобы файлы группировались по директориям
    for item in sorted(path.rglob('*'), key=lambda p: p.relative_to(base_path).as_posix()):
        if item.is_symlink():
            continue
        # Относительный путь, чтобы показать вложенность (директория "/")
        relative_path = item.relative_to(base_path).as_posix() 
        # Если мы внутри архива, добавляем префикс
        full_path = (archive_prefix + "/" + relative_path) if archive_prefix else relative_path
        
        stat = item.stat()
        mod_time = datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S')

        if item.is_file():
            size = stat.st_size
            entries.append({
                "name": full_path,
                "type": "file",
                "size": size,
                "modified": mod_time
            })
            # Если это ZIP-архив, извлекаем его содержимое во временный каталог
            if item.suffix.lower() == '.zip':
                with tempfile.TemporaryDirectory() as tmp_dir:
                    try:
                        with zipfile.ZipFile(item, 'r') as z:
                            # Определяем имя архива (относительно base_path)
                            archive_name = item.relative_to(base_path).as_posix()
                            new_archive_prefix = archive_name
                            # Перебираем имена aрхивных файлов
                            for zip_info in z.infolist():
                                # Пытаемся декодировать имя с разными кодировками
                                try:
                                    name = zip_info.filename.encode('cp437').decode('cp866')
                                except UnicodeDecodeError:
                                    try:
                                        name = zip_info.filename.encode('cp437').decode('utf-8')
                                    except UnicodeDecodeError:
                                        name = zip_info.filename  # оставляем как есть
                                
                                # Извлекаем файл с корректным именем
                                zip_info.filename = name
                                z.extract(zip_info, tmp_dir)

                                # не всегда работает детектинг
                                # raw_name = zip_info.filename.encode('cp437')  # всегда в cp437
                                # # Определяем кодировку
                                # encoding = chardet.detect(raw_name)['encoding'] or 'utf-8'
                                # # Декодируем имя
                                # decoded_name = raw_name.decode(encoding, errors='replace')
                                # # Назначаем новое имя для извлечения
                                # zip_info.filename = decoded_name
                                # z.extract(zip_info, tmp_dir)

                            # Рекурсивно сканируем извлечённое содержимое
                            entries.extend(scan_directory(Path(tmp_dir), Path(tmp_dir), archive_prefix=new_archive_prefix))
                    except zipfile.BadZipFile:
                        pass  # Игнорируем повреждённые архивы        
        elif item.is_dir():
            entries.append({
                "name": full_path,
                "type": "folder",
                "size": "0",
                "modified": mod_time
            })

    return entries

def generate_json_report(data, output_path):
    """
    Создание JSON-файла с помощью библиотеки json
    """
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def generate_csv_report(data, output_path):
    """
    Создание CSV-файла с помощью библиотеки csv
    """
    with open(output_path, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Type", "Size", "Modified"])
        for entry in data:
            writer.writerow([entry["name"], entry["type"], entry["size"], entry["modified"]])


def generate_xlsx_report(data, output_path):
    """
    Создание Excel-файла с помощью openpyxl (uv add openpyxl)
    """
    wb = Workbook()
    if wb.active is None:
        ws = wb.create_sheet("Отчёт о структуре файлов")
    else:
        ws = wb.active
    ws.title = "Отчёт о структуре файлов"
    ws.append(["Name", "Type", "Size", "Modified"])
    for entry in data:
        ws.append([entry["name"], entry["type"], entry["size"], entry["modified"]])
    wb.save(output_path)


def generate_docx_report(data, output_path):
    """
    Создание Word-файла с помощью python-docx (uv add python-docx)
    """
    doc = Document()
    doc.add_heading('Отчёт о структуре файлов', 0)

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Size'
    hdr_cells[3].text = 'Modified'

    for entry in data:
        row_cells = table.add_row().cells
        row_cells[0].text = entry["name"]
        row_cells[1].text = entry["type"]
        row_cells[2].text = str(entry["size"])
        row_cells[3].text = entry["modified"]

    doc.save(output_path)


def generate_pdf_report(data, output_path):
    """
    Создание PDF-файла с помощью reportlab (uv add reportlab)
    """
    # Регистрация шрифта Arial, чтобы не было квадратиков вместо букв
    font_path = r"C:/Windows/Fonts/arial.ttf"
    pdfmetrics.registerFont(TTFont('Arial', font_path))
    # Документ
    doc = SimpleDocTemplate(str(output_path), pagesize=A4)
    # Элементы документа
    elements = []
    # styles = getSampleStyleSheet() стиль по умолчанию - не используется
    # Создаём стиль для заголовка с Arial
    style_title = ParagraphStyle(
        name='Title',
        fontName='Arial',  # ← используем Arial
        fontSize=16,
        leading=20,
        alignment=1  # по центру
    )
    # Заголовок
    title = Paragraph("Отчёт о структуре файлов", style_title)
    elements.append(title)
    elements.append(Spacer(1, 12)) # отступ

    # Заголовки таблицы
    header_data = [["Name", "Type", "Size", "Modified"]]
    table_data = header_data + [
        [entry["name"], entry["type"], str(entry["size"]), entry["modified"]] 
        for entry in data]
    # Создаём таблицу
    table = Table(table_data)
    table.setStyle(TableStyle([
        # ('СТИЛЬ', (начало_столбца, начало_строки), (конец_столбца, конец_строки), значение)
        # 0 - первый, -1 - последний
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey), # серый фон заголовка
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke), # белый цвет текста заголовка
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'), # ← выравнивание по левому краю всей таблицы
        ('FONTNAME', (0, 0), (-1, 0), 'Arial'),  # ← шрифт для заголовка
        ('FONTNAME', (0, 1), (-1, -1), 'Arial'),  # ← шрифт для данных
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12), # ← Отступ снизу для заголовка — 12 пунктов.
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige), # ← бежевый фон для данных
        ('GRID', (0, 0), (-1, -1), 1, colors.black) # ← Сетка
    ]))

    elements.append(table)
    doc.build(elements)